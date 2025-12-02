import os
import subprocess
import shutil
from pathlib import Path
from docx import Document
from LLMClients.clients import Model

# === Step 1: find resume ending with "Resume.docx" ===
def get_resume_from_folder(folder_path):
    for file in os.listdir(folder_path):
        if file.endswith("Resume.docx"):
            return os.path.join(folder_path, file)
    raise FileNotFoundError("No resume file ending with 'Resume.docx' found.")


def convert_docx_to_pdf(docx_path: str) -> str:
    """
    Convert a DOCX file to PDF in the same folder.
    Tries methods in order: WPS Office (via AppleScript), LibreOffice, docx2pdf.
    Returns the path to the PDF file if successful, None otherwise.
    """
    docx_path_obj = Path(docx_path)
    if not docx_path_obj.exists():
        print(f"⚠️ DOCX file not found: {docx_path}")
        return None
    
    # Output PDF path (same folder, same name, .pdf extension)
    pdf_path = docx_path_obj.with_suffix('.pdf')
    output_dir = docx_path_obj.parent
    docx_abs_path = str(docx_path_obj.resolve())
    pdf_filename = pdf_path.name
    
    # Method 1: Try WPS Office (user's preferred office suite) via AppleScript
    wps_app_paths = [
        "/Applications/WPS Office.app",
        "/Applications/Kingsoft Office.app",
    ]
    
    wps_found = False
    for wps_app in wps_app_paths:
        if Path(wps_app).exists():
            wps_found = True
            try:
                # AppleScript to automate WPS Office to export to PDF
                applescript = f'''
                tell application "WPS Office"
                    activate
                    open POSIX file "{docx_abs_path}"
                end tell
                delay 3
                tell application "System Events"
                    tell process "WPS Office"
                        try
                            -- Use File > Export > Export as PDF menu
                            click menu bar item "File" of menu bar 1
                            delay 0.5
                            click menu item "Export" of menu "File" of menu bar 1
                            delay 0.5
                            click menu item "Export as PDF" of menu "Export" of menu item "Export" of menu "File" of menu bar 1
                            delay 2
                            -- Save dialog - enter filename and save
                            keystroke "{pdf_filename}"
                            delay 0.5
                            key code 36 -- Return key
                            delay 2
                        on error errMsg
                            return "Error: " & errMsg
                        end try
                    end tell
                end tell
                delay 2
                tell application "WPS Office"
                    quit
                end tell
                '''
                
                result = subprocess.run(
                    ["osascript", "-e", applescript],
                    capture_output=True,
                    text=True,
                    timeout=30
                )
                
                if result.returncode == 0 and pdf_path.exists():
                    print(f"✅ PDF exported using WPS Office: {pdf_path}")
                    return str(pdf_path)
                elif pdf_path.exists():
                    print(f"✅ PDF exported using WPS Office: {pdf_path}")
                    return str(pdf_path)
                else:
                    print(f"⚠️ WPS Office automation attempted but PDF not created")
            except (subprocess.TimeoutExpired, FileNotFoundError, Exception) as e:
                print(f"⚠️ WPS Office conversion failed: {e}")
            break
    
    # Method 2: Try LibreOffice (most reliable command-line method)
    libreoffice_paths = [
        "/Applications/LibreOffice.app/Contents/MacOS/soffice",
        "/usr/local/bin/soffice",
        "/opt/homebrew/bin/soffice",
        shutil.which("soffice"),
    ]
    
    libreoffice_cmd = None
    for path in libreoffice_paths:
        if path and Path(path).exists():
            libreoffice_cmd = path
            break
    
    if libreoffice_cmd:
        try:
            # LibreOffice command to convert to PDF
            cmd = [
                libreoffice_cmd,
                "--headless",
                "--convert-to", "pdf",
                "--outdir", str(output_dir),
                str(docx_path_obj)
            ]
            result = subprocess.run(
                cmd,
                capture_output=True,
                text=True,
                timeout=60
            )
            
            if result.returncode == 0 and pdf_path.exists():
                print(f"✅ PDF exported using LibreOffice: {pdf_path}")
                return str(pdf_path)
        except (subprocess.TimeoutExpired, FileNotFoundError, Exception) as e:
            print(f"⚠️ LibreOffice conversion failed: {e}")
    
    # Method 3: Try docx2pdf (works with Microsoft Word, may work with WPS)
    try:
        from docx2pdf import convert
        convert(docx_path, str(pdf_path))
        
        if pdf_path.exists():
            print(f"✅ PDF exported using docx2pdf: {pdf_path}")
            return str(pdf_path)
        else:
            print(f"⚠️ PDF conversion completed but file not found: {pdf_path}")
    except ImportError:
        if not wps_found:
            print("⚠️ docx2pdf not installed. Install it with: pip install docx2pdf")
    except Exception as e:
        error_msg = str(e)
        if "Message not understood" in error_msg or "not found" in error_msg.lower():
            if wps_found:
                print(f"⚠️ PDF conversion failed: Could not automate WPS Office.")
                print(f"💡 Tip: Try installing LibreOffice for reliable command-line PDF conversion: brew install --cask libreoffice")
            else:
                print(f"⚠️ PDF conversion failed: Office application not available.")
                print(f"💡 Tip: Install LibreOffice for reliable command-line PDF conversion: brew install --cask libreoffice")
        else:
            print(f"⚠️ PDF conversion failed: {error_msg}")
    
    return None

# === Step 2: get cover letter template path ===
cover_letter_path_file = "/Users/Roger/Documents/FullTime-Resume/Resume Template - One Page/cover_letter_path.txt"
with open(cover_letter_path_file, "r") as f:
    cover_letter_template_path = f.read().strip()

resume_folder = os.path.dirname(cover_letter_template_path)
resume_path = get_resume_from_folder(resume_folder)

# === Step 3: read job description ===
jd_path = os.path.join(os.path.dirname(cover_letter_template_path), "job_description.txt")
with open(jd_path, "r") as f:
    job_description = f.read()

# === Step 4: prepare prompt ===
with open(resume_path, "rb") as f:
    resume_doc = Document(f)
    resume_text = "\n".join([p.text for p in resume_doc.paragraphs])

prompt = f"""
You are an assistant that writes professional job application cover letters.
Use the resume and job description below to generate only the **body text**
of the cover letter (no greeting like "Hi Hiring Team" and no ending with my name).
Should be confident and professional, and a bit natural, but not using "—" or "'", and not too much focus on the things repetitively mentioned in resume, i meant not too much detail and no less than 250 and no more than 300 words.

Resume:
{resume_text}

Job Description:
{job_description}
"""

llm = Model("GEMINI", prompt)

# # response = client.chat.completions.create(
# #     model="gpt-5-mini",   # you can adjust model
# #     messages=[{"role": "user", "content": prompt}],
# # )
# response = client.models.generate_content(
#     model="gemini-2.5-pro", contents=prompt
# )

# # cover_letter_body = response.choices[0].message.content.strip()
# cover_letter_body = response.candidates[0].content.parts[0].text
cover_letter_body = llm.get_response_from_client()

# === Step 5: insert into template ===
doc = Document(cover_letter_template_path)

for para in doc.paragraphs:
    if "{{COVER_LETTER_BODY}}" in para.text:
        para.text = cover_letter_body
        para.style = "NewCoverLetterStyle"  # Apply your custom style
        break

output_path = cover_letter_template_path.replace("_Template.docx", ".docx")
doc.save(output_path)

print(f"✅ Cover letter generated: {output_path}")

# Export cover letter as PDF
print("📄 Exporting cover letter to PDF...")
convert_docx_to_pdf(output_path)
