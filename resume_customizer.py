# resume_customizer.py
# from openai import OpenAI
import os, json
from docx import Document
# from dotenv import load_dotenv
import re
from docx.oxml import OxmlElement
from LLMClients.clients import Model
from job_description_cleaner.jd_cleaning import clean_job_description
from resume_prompts import SECTION_STYLES, build_prompt, get_sections_for_category
# ---------- CONFIG ----------
# load_dotenv()
# client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
# MODEL = "gpt-5-mini"   # or "gpt-5-mini"

# ---------- LLM HELPER ----------
import time
from typing import Dict, Optional

from openai import APIError, RateLimitError

MAX_LLM_RETRIES = int(os.getenv("LLM_MAX_RETRIES", "4"))
BACKOFF_BASE_SECONDS = float(os.getenv("LLM_BACKOFF_BASE", "1.5"))


def _call_llm_with_retries(prompt: str) -> Optional[str]:
    for attempt in range(MAX_LLM_RETRIES + 1):
        try:
            client = Model("OPENAI", prompt)
            response = client.get_response_from_client()
            return response
        except (RateLimitError, APIError) as exc:
            if attempt == MAX_LLM_RETRIES:
                print(f"⚠️ LLM call failed after {MAX_LLM_RETRIES + 1} attempts: {exc}")
                return None
            sleep_seconds = BACKOFF_BASE_SECONDS * (2 ** attempt)
            print(f"⚠️ LLM call failed (attempt {attempt + 1}/{MAX_LLM_RETRIES + 1}): {exc}. Retrying in {sleep_seconds:.1f}s...")
            time.sleep(sleep_seconds)
        except Exception as exc:
            print(f"⚠️ Unexpected LLM error: {exc}")
            return None


def improve_resume_json(section_texts: Dict[str, str], job_description: str, job_category: str = "sde", additional_info: Optional[str] = None) -> Dict[str, str]:
    """Ask LLM to improve all resume sections in one JSON response."""
    # Build the prompt using the prompt builder based on job category
    prompt = build_prompt(section_texts, job_description, job_category, additional_info)
    
    response = _call_llm_with_retries(prompt)
    if not response:
        print("⚠️ Falling back to original sections due to LLM failure.")
        # Return only the sections that should be included for this category
        sections_to_include = get_sections_for_category(job_category)
        return {k: v for k, v in section_texts.items() if k in sections_to_include}
    try:
        improved = json.loads(response)
        # Ensure we only return sections that should be included for this category
        sections_to_include = get_sections_for_category(job_category)
        return {k: v for k, v in improved.items() if k in sections_to_include}
    except json.JSONDecodeError:
        print("⚠️ Model did not return valid JSON, raw output saved for debugging")
        return {}

# ---------- MAIN PIPELINE ----------
def add_text_with_bold(para, text):
    """
    Add text to a paragraph, bolding any content wrapped in **...**.
    """
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = para.add_run(part[2:-2])  # strip **
            run.bold = True
        else:
            para.add_run(part)
            
def insert_paragraph_after(paragraph, style=None):
    """Insert a new paragraph immediately after the given one."""
    new_para = paragraph._parent.add_paragraph("", style=style)  # create paragraph
    paragraph._element.addnext(new_para._element)  # insert directly after
    return new_para

def customize_resume_with_placeholders(template_path: str, section_files: dict, job_description: str, output_path: str, job_category: str = "sde", additional_info: str = None):
    """
    Replace placeholders in resume template with LLM-customized content.
    section_files = {"SUMMARY": "path/to/summary.txt", ...}
    job_category: "sde" or "support" - determines which sections and prompts to use
    """
    doc = Document(template_path)

    # Load all base texts
    section_texts = {}
    for placeholder, file_path in section_files.items():
        if os.path.exists(file_path):
            with open(file_path, "r") as f:
                section_texts[placeholder] = f.read()
        else:
            print(f"⚠️ Section file not found: {file_path}")

    # Clean job description
    cleaned_job_description = clean_job_description(job_description)

    # Get improved sections from LLM (filtered by job category)
    improved_sections = improve_resume_json(section_texts, cleaned_job_description, job_category, additional_info)
    # print the text of each section for debugging
    for section, text in improved_sections.items():
        print(f"--- {section} ---\n{text}\n")   
    if not improved_sections:
        print("❌ No improved sections returned, aborting placeholder replacement.")
        return

    # Build replacements dict ({{SUMMARY}} → improved text, etc.)
    replacements = {f"{{{{{k}}}}}": v for k, v in improved_sections.items()}

    # Replace in paragraphs and tables
    def replace_placeholders_in_doc(doc_obj, replacements: dict):
        for para in doc_obj.paragraphs:
            for key, value in replacements.items():
                if key in para.text:
                    section_name = key.strip("{}")   # e.g. "SUMMARY"
                    style = SECTION_STYLES.get(section_name, "Normal")

                    # split into lines (ignore empty lines)
                    items = [line.strip() for line in value.splitlines() if line.strip()]

                    # clear the placeholder paragraph
                    para.text = ""

                    # add first item
                    if items:
                        para.style = style
                        add_text_with_bold(para, items[0])

                        # insert the rest as new styled paragraphs (keep correct order)
                        current_para = para
                        for item in items[1:]:
                            new_para = insert_paragraph_after(current_para, style=style)
                            add_text_with_bold(new_para, item)
                            current_para = new_para  # advance pointer

        # Handle table placeholders too
        for table in doc_obj.tables:
            for row in table.rows:
                for cell in row.cells:
                    replace_placeholders_in_doc(cell, replacements)

    replace_placeholders_in_doc(doc, replacements)

    doc.save(output_path)
    print(f"✅ Customized resume saved to {output_path}")
 
